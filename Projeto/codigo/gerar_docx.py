# -*- coding: utf-8 -*-
"""
gerar_docx.py - Gera o relatorio final em .docx usando o template oficial da Vale como base.

Abre `Original/Desenvolver_Template.docx` (herdando cabecalho com logo, estilos e
layout), limpa o corpo de exemplo e reconstroi o documento a partir de
`Projeto/relatorio/relatorio_final.md`: pagina de rosto, pagina de autor, resumo,
4 secoes, tabela e figuras embutidas.

Saida (arquivo de entrega final, movido para a raiz do repositorio em 19/07):
    Relatorio_Final_Marcelo_Ayala_Gomes.docx

Executar:
    PYTHONIOENCODING=utf-8 uv run python Projeto/codigo/gerar_docx.py
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

try:
    from PIL import Image
    HAS_PIL = True
except Exception:
    HAS_PIL = False

ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = ROOT / "Original" / "Desenvolver_Template.docx"
MD = ROOT / "Projeto" / "relatorio" / "relatorio_final.md"
FIG_DIR = ROOT / "Projeto" / "relatorio"
OUT = ROOT / "Relatorio_Final_Marcelo_Ayala_Gomes.docx"

CENTER = WD_ALIGN_PARAGRAPH.CENTER
JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY
# recuo deslocado (hanging) das listas, igual a versao revisada pelo autor
LIST_LEFT = Inches(0.5)
LIST_HANG = Inches(-0.25)


def clear_body(doc):
    """Remove paragrafos e tabelas do corpo, mantendo o sectPr no lugar (necessario para add_table)."""
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child.tag == qn("w:p") or child.tag == qn("w:tbl"):
            body.remove(child)
    return body, sectPr


def clean_header(doc):
    """Remove o texto 'Template Padrao' do cabecalho, preservando o logo (runs com desenho)."""
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            for r in p.runs:
                tem_desenho = (r._element.find(qn("w:drawing")) is not None
                               or bool(r._element.findall(".//" + qn("a:blip"))))
                if not tem_desenho and r.text:
                    r.text = ""


def style_names(doc):
    names = set()
    for s in doc.styles:
        try:
            names.add(s.name)
        except Exception:
            pass
    return names


INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)")


def add_runs(p, text):
    """Adiciona runs interpretando **negrito**, *italico* e `codigo`.
    Texto normal herda o tamanho do template (12pt uniforme, docDefaults); codigo em 9pt."""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        tok = m.group()
        if tok.startswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        elif tok.startswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def add_image(doc, caption, relpath):
    # legenda ACIMA da figura (justificada, italico 9pt), imagem centralizada
    cap = doc.add_paragraph()
    cap.alignment = JUSTIFY
    rc = cap.add_run(caption)
    rc.italic = True
    rc.font.size = Pt(9)

    img = FIG_DIR / relpath
    if not img.exists():
        print(f"  [AVISO] figura ausente, pulada: {relpath}")
        p = doc.add_paragraph()
        p.alignment = CENTER
        add_runs(p, "[figura ausente]")
        return
    wc = 15.0
    if HAS_PIL:
        try:
            w, h = Image.open(img).size
            aspect = h / w
            if wc * aspect > 20.0:
                wc = 20.0 / aspect
        except Exception:
            pass
    doc.add_picture(str(img), width=Cm(wc))
    doc.paragraphs[-1].alignment = CENTER


def build_title_pages(doc):
    def centered(text, size, bold=True, italic=False, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(space_before); pf.space_after = Pt(space_after)
        r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        return p

    # Pagina de rosto
    for _ in range(6):
        doc.add_paragraph()
    centered("Relatório Final", 24)
    centered("Desafio: Análise Avançada de Dados", 16)
    doc.add_paragraph()
    centered("Antecipação de Alertas Don't Go em Frotas de Mineração", 14, italic=True)
    doc.add_page_break()

    # Pagina de autor
    for _ in range(6):
        doc.add_paragraph()
    centered("Marcelo Henrique Ayala Gomes", 13)
    centered("UNESP Bauru", 12, bold=False)
    centered("mh.gomes@unesp.br", 12, bold=False)
    doc.add_paragraph()
    centered("Trabalho individual (sem grupo)", 10, bold=False, italic=True)
    doc.add_paragraph()
    centered("Programa Desenvolver 2026, Vale. Região: Itabira (MG). Dados: janeiro a junho de 2025.",
             10, bold=False)
    centered("Código e materiais completos: https://github.com/MarceloHAyala/DadosVale.git",
             10, bold=False)
    doc.add_page_break()


def parse_table(lines, i, doc, styles):
    # coleta linhas consecutivas que comecam com '|'
    rows = []
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        rows.append(lines[i])
        i += 1
    # remove linha separadora (---)
    parsed = []
    for r in rows:
        cells = [c.strip() for c in r.strip().strip("|").split("|")]
        if all(set(c) <= set("-: ") for c in cells):
            continue
        parsed.append(cells)
    if not parsed:
        return i
    ncols = len(parsed[0])
    table = doc.add_table(rows=len(parsed), cols=ncols)
    if "Table Grid" in styles:
        table.style = "Table Grid"
    for ri, row in enumerate(parsed):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            txt = row[ci] if ci < len(row) else ""
            add_runs(p, txt)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()
    return i


def build_body(doc, md_text, styles):
    lines = md_text.split("\n")
    i = 0
    heading2 = "Heading 2" if "Heading 2" in styles else ("Heading 1" if "Heading 1" in styles else "Normal")
    list_num = "List Number" if "List Number" in styles else ("List Bullet" if "List Bullet" in styles else "Normal")
    list_bul = "List Bullet" if "List Bullet" in styles else "Normal"

    while i < len(lines):
        line = lines[i].rstrip()
        s = line.strip()

        if s == "" or s == "---":
            i += 1
            continue

        # tabela
        if s.startswith("|"):
            i = parse_table(lines, i, doc, styles)
            continue

        # figura: [texto](figuras/...)
        mfig = re.match(r"^\[(.+?)\]\((figuras/.+?)\)$", s)
        if mfig:
            add_image(doc, mfig.group(1), mfig.group(2))
            i += 1
            continue

        # headings (todos sem numero; Resumo e as 4 secoes -> Heading 1; subsecoes -> Heading 2)
        if s.startswith("## ") and not s.startswith("### "):
            h = doc.add_paragraph(s[3:].strip(),
                                  style="Heading 1" if "Heading 1" in styles else "Normal")
            h.alignment = JUSTIFY
            i += 1
            continue
        if s.startswith("### "):
            h = doc.add_paragraph(s[4:].strip(), style=heading2)
            h.alignment = JUSTIFY
            i += 1
            continue
        if s.startswith("#### "):
            p = doc.add_paragraph(); p.alignment = JUSTIFY
            r = p.add_run(s[5:].strip()); r.bold = True
            i += 1
            continue

        # blockquote (pergunta central): justificado, negrito + italico, recuo a esquerda
        if s.startswith("> "):
            p = doc.add_paragraph()
            p.alignment = JUSTIFY
            p.paragraph_format.left_indent = Cm(1.25)
            add_runs(p, s[2:].strip())
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        # lista numerada -> paragrafo normal justificado com recuo deslocado, numero removido
        mnum = re.match(r"^\d+\.\s+(.*)$", s)
        if mnum:
            p = doc.add_paragraph(); p.alignment = JUSTIFY
            p.paragraph_format.left_indent = LIST_LEFT
            p.paragraph_format.first_line_indent = LIST_HANG
            add_runs(p, mnum.group(1))
            i += 1
            continue

        # lista com marcador -> paragrafo normal justificado com recuo deslocado, marcador removido
        if s.startswith("- "):
            p = doc.add_paragraph(); p.alignment = JUSTIFY
            p.paragraph_format.left_indent = LIST_LEFT
            p.paragraph_format.first_line_indent = LIST_HANG
            add_runs(p, s[2:].strip())
            i += 1
            continue

        # paragrafo normal justificado
        p = doc.add_paragraph(); p.alignment = JUSTIFY
        add_runs(p, s)
        i += 1

    return doc


def main():
    print("Abrindo template:", TEMPLATE.name)
    doc = Document(str(TEMPLATE))
    styles = style_names(doc)
    print("Estilos disponiveis (relevantes):",
          [x for x in ["Normal", "Heading 1", "Heading 2", "List Bullet", "List Number", "Table Grid"] if x in styles])

    clean_header(doc)  # remove "Template Padrao" do cabecalho, preserva o logo
    body, sectPr = clear_body(doc)

    build_title_pages(doc)

    md_text = MD.read_text(encoding="utf-8")
    # pega do "## Resumo" em diante
    idx = md_text.find("## Resumo")
    body_md = md_text[idx:] if idx >= 0 else md_text
    build_body(doc, body_md, styles)

    # move o sectPr para o final do corpo (ordem valida do documento)
    if sectPr is not None:
        doc.element.body.remove(sectPr)
        doc.element.body.append(sectPr)

    out_path = Path(sys.argv[1]) if len(sys.argv) > 1 else OUT
    doc.save(str(out_path))
    print("Salvo:", out_path)
    print(f"Paragrafos: {len(doc.paragraphs)} | Tabelas: {len(doc.tables)}")


if __name__ == "__main__":
    main()
